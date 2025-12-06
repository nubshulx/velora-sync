"""
Word document reader for extracting requirements
"""

from pathlib import Path
from typing import Optional, List, Dict
import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
import tempfile

from src.utils.exceptions import DocumentReadError
from src.utils.logger import get_logger
from src.document_readers.sharepoint_client import SharePointClient

logger = get_logger(__name__)


class WordReader:
    """Reader for Word documents containing requirements"""
    
    def __init__(
        self,
        sharepoint_client: Optional[SharePointClient] = None
    ):
        """
        Initialize Word reader
        
        Args:
            sharepoint_client: Optional SharePoint client for cloud documents
        """
        self.sharepoint_client = sharepoint_client
    
    def read_document(self, document_path: str) -> str:
        """
        Read requirements from Word document
        
        Args:
            document_path: Path to Word document (local or SharePoint URL)
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentReadError: If reading fails
        """
        try:
            # Determine if path is SharePoint URL
            is_sharepoint = 'sharepoint.com' in document_path.lower()
            
            if is_sharepoint:
                if not self.sharepoint_client:
                    raise DocumentReadError("SharePoint client not configured")
                
                # Download to temporary file
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                    tmp_path = Path(tmp_file.name)
                
                logger.info(f"Downloading Word document from SharePoint: {document_path}")
                self.sharepoint_client.download_file(document_path, tmp_path)
                
                try:
                    content = self._extract_text(tmp_path)
                finally:
                    # Clean up temporary file
                    if tmp_path.exists():
                        tmp_path.unlink()
            else:
                # Read local file
                local_path = Path(document_path)
                if not local_path.exists():
                    raise DocumentReadError(f"Document not found: {document_path}")
                
                logger.info(f"Reading local Word document: {document_path}")
                content = self._extract_text(local_path)
            
            logger.info(f"Successfully extracted {len(content)} characters from document")
            return content
            
        except DocumentReadError:
            raise
        except Exception as e:
            logger.error(f"Failed to read Word document: {str(e)}")
            raise DocumentReadError(f"Failed to read document: {str(e)}")
    
    def _extract_text(self, file_path: Path) -> str:
        """
        Extract text content from Word document
        
        Args:
            file_path: Path to local Word document
            
        Returns:
            Extracted text with structure preserved
        """
        try:
            doc: Document = docx.Document(str(file_path))
            
            # Extract text with structure
            content_parts: List[str] = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Preserve heading structure
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.replace('Heading ', '')
                        content_parts.append(f"\n{'#' * int(level) if level.isdigit() else '#'} {text}\n")
                    else:
                        content_parts.append(text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        content_parts.append(row_text)
            
            return '\n'.join(content_parts)
            
        except Exception as e:
            raise DocumentReadError(f"Failed to extract text from document: {str(e)}")
    
    def extract_requirements(self, document_path: str) -> List[Dict[str, str]]:
        """
        Extract structured requirements from Word document
        Handles both well-formatted and unstructured documents
        
        Args:
            document_path: Path to Word document
            
        Returns:
            List of requirement dictionaries with 'id', 'title', and 'content'
        """
        try:
            content = self.read_document(document_path)
            
            # Parse requirements from content
            requirements = []
            current_requirement = None
            
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Detect requirement headers (e.g., "REQ-001:", "Requirement 1:", etc.)
                if self._is_requirement_header(line):
                    # Save previous requirement
                    if current_requirement:
                        requirements.append(current_requirement)
                    
                    # Start new requirement
                    req_id, title = self._parse_requirement_header(line)
                    current_requirement = {
                        'id': req_id,
                        'title': title,
                        'content': []
                    }
                elif current_requirement is not None:
                    # Add to current requirement content
                    current_requirement['content'].append(line)
            
            # Add last requirement
            if current_requirement:
                requirements.append(current_requirement)
            
            # Join content lines
            for req in requirements:
                req['content'] = '\n'.join(req['content'])
            
            # If no structured requirements found, try intelligent chunking
            if not requirements:
                logger.warning("No structured requirements found, attempting intelligent chunking")
                requirements = self._chunk_unstructured_content(content)
            
            logger.info(f"Extracted {len(requirements)} requirement(s)")
            return requirements
            
        except Exception as e:
            logger.error(f"Failed to extract requirements: {str(e)}")
            # Last resort: return entire content as single requirement
            content = self.read_document(document_path)
            return [{
                'id': 'REQ-001',
                'title': 'General Requirements',
                'content': content
            }]
    
    def _chunk_unstructured_content(self, content: str) -> List[Dict[str, str]]:
        """
        Intelligently chunk unstructured content into logical requirements
        
        Args:
            content: Full document content
            
        Returns:
            List of requirement dictionaries
        """
        requirements = []
        
        # Strategy 1: Split by markdown-style headings (# Heading)
        heading_pattern = r'^#+\s+(.+)$'
        import re
        
        sections = []
        current_section = {'title': 'Introduction', 'content': []}
        
        for line in content.split('\n'):
            match = re.match(heading_pattern, line.strip())
            if match:
                # Save previous section
                if current_section['content']:
                    sections.append(current_section)
                # Start new section
                current_section = {
                    'title': match.group(1).strip(),
                    'content': []
                }
            else:
                if line.strip():
                    current_section['content'].append(line)
        
        # Add last section
        if current_section['content']:
            sections.append(current_section)
        
        # If we found sections, use them as requirements
        if len(sections) > 1:
            for idx, section in enumerate(sections, 1):
                requirements.append({
                    'id': f'REQ-{idx:03d}',
                    'title': section['title'],
                    'content': '\n'.join(section['content'])
                })
            logger.info(f"Chunked by headings: {len(requirements)} sections")
            return requirements
        
        # Strategy 2: Split by double line breaks (paragraph-based)
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        if len(paragraphs) > 1 and len(paragraphs) <= 50:
            # Reasonable number of paragraphs
            for idx, para in enumerate(paragraphs, 1):
                # Use first 50 chars as title
                title = para[:50] + '...' if len(para) > 50 else para
                requirements.append({
                    'id': f'REQ-{idx:03d}',
                    'title': title,
                    'content': para
                })
            logger.info(f"Chunked by paragraphs: {len(requirements)} requirements")
            return requirements
        
        # Strategy 3: Split by character count (for very large unstructured docs)
        if len(content) > 5000:
            chunk_size = 2000  # ~2000 chars per requirement
            chunks = []
            
            for i in range(0, len(content), chunk_size):
                chunk = content[i:i + chunk_size]
                # Try to break at sentence boundary
                if i + chunk_size < len(content):
                    last_period = chunk.rfind('.')
                    if last_period > chunk_size * 0.7:  # At least 70% through
                        chunk = chunk[:last_period + 1]
                
                chunks.append(chunk.strip())
            
            for idx, chunk in enumerate(chunks, 1):
                title = f"Requirement Section {idx}"
                requirements.append({
                    'id': f'REQ-{idx:03d}',
                    'title': title,
                    'content': chunk
                })
            logger.info(f"Chunked by size: {len(requirements)} chunks")
            return requirements
        
        # Strategy 4: Treat entire document as single requirement
        logger.info("Using entire document as single requirement")
        return [{
            'id': 'REQ-001',
            'title': 'Complete Requirements Document',
            'content': content
        }]
    
    def _is_requirement_header(self, line: str) -> bool:
        """Check if line is a requirement header"""
        line_lower = line.lower()
        
        # Common requirement header patterns
        patterns = [
            'req-', 'req_', 'requirement',
            'fr-', 'fr_', 'functional requirement',
            'nfr-', 'nfr_', 'non-functional requirement',
            'user story', 'us-', 'us_'
        ]
        
        return any(pattern in line_lower for pattern in patterns)
    
    def _parse_requirement_header(self, line: str) -> tuple[str, str]:
        """
        Parse requirement header to extract ID and title
        
        Args:
            line: Header line
            
        Returns:
            Tuple of (requirement_id, title)
        """
        # Try to split on common delimiters
        for delimiter in [':', '-', '|']:
            if delimiter in line:
                parts = line.split(delimiter, 1)
                req_id = parts[0].strip()
                title = parts[1].strip() if len(parts) > 1 else ''
                return req_id, title
        
        # If no delimiter, use whole line as title
        return 'REQ', line.strip()
